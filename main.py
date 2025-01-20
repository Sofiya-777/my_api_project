from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel  # Импортируем BaseModel для обработки данных
from docx import Document

# Создаём приложение FastAPI
app = FastAPI()

# Определяем модель данных для входящего запроса
class DocumentData(BaseModel):
    title: str
    content: str

@app.get("/")
def read_root():
    """Приветственное сообщение"""
    return {"message": "Добро пожаловать! Это ваш API для генерации документов."}

@app.post("/generate")
def generate_document(data: DocumentData):
    """Генерация документа Word"""
    # Извлекаем данные из тела запроса
    title = data.title
    content = data.content

    # Создаём Word-документ
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

    # Сохраняем документ в файл
    filename = f"{title}.docx"
    doc.save(filename)

    # Отправляем документ пользователю
    return FileResponse(filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)
